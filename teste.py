from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from openpyxl import load_workbook

import win32com.client as win32


outlook = win32.Dispatch("outlook.application")


nome_arquivo = "Empresas.xlsx"
nome_planilha = load_workbook(nome_arquivo)
nome_guia = nome_planilha["Relação"]

for linha in range(2, len(nome_guia["A"]) + 1):

    arquivo = Document("Modelo1.docx")

    estilo = arquivo.styles["Normal"]

    pessoa = nome_guia['A%s' % linha].value
    dia = nome_guia['B%s' % linha].value
    mes = nome_guia['C%s' % linha].value
    ano = nome_guia['D%s' % linha].value
    curso = nome_guia['E%s' % linha].value
    instrutor = nome_guia['F%s' % linha].value
    emailAlvo = nome_guia['G%s' % linha].value

    descricao1 = f"Completou o curso de "
    descricao2 = f" em {dia} de {mes} de {ano}."

    for paragrafo in arquivo.paragraphs:
        if "@nome" in paragrafo.text:
            paragrafo.text = pessoa
            fonte = estilo.font
            fonte.name = "Calabri (Corpo)"
            fonte.size = Pt(12)

        if "@descricao" in paragrafo.text:
            paragrafo.text = descricao1
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(12)
            adicionarNovaPalavra = paragrafo.add_run(curso)
            adicionarNovaPalavra.font.color.rgb = RGBColor(255, 0, 0)
            adicionarNovaPalavra.underline = True
            adicionarNovaPalavra.bold = True
            adicionarNovaPalavra = paragrafo.add_run(descricao2)
            adicionarNovaPalavra.font.color.rgb = RGBColor(0, 0, 0)
            adicionarNovaPalavra.underline = False
            adicionarNovaPalavra.bold = False

        if "@instrutor" in paragrafo.text:
            paragrafo.text = instrutor + " (Tutor)"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(12)

    localDocumento = "gerados/" + pessoa + ".docx"
    arquivo.save(localDocumento)

    emailOutlook = outlook.CreateItem(0)
    emailOutlook.To = emailAlvo
    emailOutlook.Subject = "Seu certificado chegou " + pessoa + "!!!"
    emailOutlook.HTMLBody = f"""
        <p>Bom dia {pessoa},</p>
        <p>segue em anexo o seu <b>certificado</b>.</p>
        <p>Atenciosamente,</p>
        <p>{instrutor}</p>
    """
    emailOutlook.Attachments.Add(localDocumento)
    emailOutlook.save()

print("Gerado com sucesso!")
